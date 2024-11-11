import os
import io
from io import BytesIO
import re
import tempfile
from typing import List, Literal, Union
import warnings
import bleach


from django.template.loader import get_template
from django.shortcuts import render, redirect
from django.template.loader import get_template
from django.http import FileResponse, HttpResponse
from django.db.models import Q
from django.utils.html import strip_tags
from datetime import date
from .models import Year
from .forms import *


from PyPDF2 import PdfMerger, PdfReader, PdfWriter

from reportlab.lib.units import cm
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import BaseDocTemplate, PageTemplate, KeepTogether
from reportlab.platypus import Frame, PageTemplate, Spacer
from reportlab.lib.units import cm
from reportlab.platypus import BaseDocTemplate
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
# Create your views here.
from xhtml2pdf import pisa





def create_report(request, pk):
    responsible_ministry_name = UserSector.objects.get(user=request.user.id).user_sector.responsible_ministry_eng
    report  = Report.objects.filter(report_type__id =pk, responsible_ministry__responsible_ministry_eng = responsible_ministry_name ).first()
    if report is not None:
        url = reverse('update_report', kwargs={'pk': report.id})
        return redirect(url)
    reports = list(Report.objects.values_list('report_type', flat=True))
    rep_type = ReportType.objects.get(Q(deadline__gte=datetime.date.today(), id = pk))


    form = ReportForm()
    if request.method == 'POST':
        form = ReportForm(request.POST, request.FILES)
        if form.is_valid():
            form.instance.report_type = rep_type
            recorded_by = Account.objects.get(id=request.user.id)
            form.instance.recorded_by = recorded_by
            form.instance.responsible_ministry = UserSector.objects.get(user=request.user.id).user_sector
            form.instance.quarter = rep_type.quarter
            form.instance.year = rep_type.year
        

            form.save()

            report = Report.objects.get(id=form.instance.id)

            for section in rep_type.report_guideline.all():
                title = section.section_title
                report_section = ReportSection.objects.create(title=title, section_content=request.POST.get('_content'), recorded_by=form.instance.recorded_by, report =report )
                report_section.save()
                # report.report_section.add(report_section)
            return redirect('view_quarter_reports')

    context = {
        'form': form, 'form_button': 'Create', 'form_method': 'POST', 'form_title': 'Create report','filefield': True, 
        'page_title': 'Create report', 'toolbar_url': 'view_quarter_reports', 'toolbar_url_name': 'Cancel', 'back': True,
        'responsible_ministry_name': responsible_ministry_name, 'rep_type': rep_type,
        }
    return render(request, 'report_form.html', context)





def list_master_reports(request):
    report_type = ReportType.objects.latest('id')
    deadline = report_type.deadline
    deadline_status = deadline < datetime.date.today()
    
    master_reports = Report.objects.filter(responsible_ministry = UserSector.objects.get(user=request.user.id).user_sector)
    context = {'master_reports': master_reports, 
               'page_title': 'Master reports', 
               'toolbar_url': 'create-report', 
               'toolbar_url_name': 'Create master report',
                'deadline': deadline_status,
                'back': True}
    return render(request, 'report.html', context)


def grid_quarter_report(request):

    report_type = ReportType.objects.order_by('-deadline').first()
    
    current_year = date.today().year  + 2
    years = Year.objects.filter(year_eng__lte=current_year, visible=True)
    quarters = Quarter.objects.all()
    reports = Report.objects.filter(responsible_ministry = UserSector.objects.get(user=request.user.id).user_sector )

    context = {
        'years': years,
        'quarters':quarters,
        'reports':reports,
        'report_type':report_type,
        
               }
    return render(request, 'grid_quarter_reports.html', context)




def detail_report(request, pk):

    responsible_ministry_name = UserSector.objects.get(user=request.user.id).user_sector.responsible_ministry_eng
    report = Report.objects.get(id = pk)
    report_contents = ReportSection.objects.filter(report=report)
    context = {
                'report': report, 'page_title': "Detail report", 'toolbar_url': 'view_quarter_reports',
                'back': True, 'responsible_ministry_name': responsible_ministry_name,
                'report_contents': report_contents
              }

    response = render(request, 'detail_report.html', context)
    response['Content-Disposition'] = 'inline'
    return response





def update_report(request, pk):
    responsible_ministry_name = UserSector.objects.get(user=request.user.id).user_sector.responsible_ministry_eng
    report = Report.objects.get(id=pk)
    report_contents = ReportSection.objects.filter(report__id = pk)
    current_table = report.report_document
    current_table_name = os.path.basename(report.report_document.name)
    form = ReportForm(instance=report)
    if report.report_type.deadline >  datetime.date.today():
        if request.method == 'POST':
            form = ReportForm(request.POST, request.FILES, instance=report)
            
            if form.is_valid():
                for report_content in report_contents:
                    report_content_input = request.POST.get(f"report_content_{report_content.id}_content")
                    current_report_content = report_contents.get(id=report_content.id)
                    current_report_content.section_content = report_content_input
                    current_report_content.save()
                form.save()
                return redirect('detail_report', pk=report.pk)

    context = {
        'form': form, 'form_button': 'Update', 'form_method': 'POST', 'form_title': 'Update report','filefield': True, 
        'page_title': 'Update report', 'toolbar_url': 'view_quarter_reports', 'toolbar_url_name': 'Cancel', 'back': True,
        'responsible_ministry_name': responsible_ministry_name, 'report_contents': report_contents, 'current_table':current_table, 'current_table_name':current_table_name
        }
    return render(request, 'update_report.html', context)








def watermark(
    content_pdf,
    stamp_pdf,
    pdf_result,
    page_indices: Union[Literal["ALL"], List[int]] = "ALL",
):
    reader = PdfReader(content_pdf)
    if page_indices == "ALL":
        page_indices = list(range(0, len(reader.pages)))

    writer = PdfWriter()
    for index in page_indices:
        content_page = reader.pages[index]
        mediabox = content_page.mediabox

        
        reader_stamp = PdfReader(stamp_pdf)
        image_page = reader_stamp.pages[0]

        image_page.merge_page(content_page)
        image_page.mediabox = mediabox
        writer.add_page(image_page)

    with open(pdf_result, "wb") as fp:
        writer.write(fp)

# def pdf_report(request, pk):
#     """
#     Generating the PDF file for the report content
#     """
#     #Stylesheet configs
#     font_family = 'Amharic'
#     font_size = 10
#     document = generate_report_word_document(request, pk)
#     return document
    # pdfmetrics.registerFont(TTFont('AbyssinicaSIL-R', './static/fonts/NotoSerifEthiopic-Regular.ttf'))
    # pdfmetrics.registerFont(TTFont('AbyssinicaSIL-B', './static/fonts/NotoSerifEthiopic-Bold.ttf'))
    # stylesheet = getSampleStyleSheet()
    # stylesheet.add(ParagraphStyle(name='Bold', fontName='AbyssinicaSIL-B', spaceBefore=5, leading=16))
    # stylesheet.add(ParagraphStyle(name='Regular', fontName='AbyssinicaSIL-R'))

    # text_frame = Frame(
    #     x1=2.54 * cm,
    #     y1=4 * cm,  
    #     height=20.85 * cm,
    #     width=16.95 * cm,
    #     leftPadding=0 * cm,
    #     bottomPadding=0 * cm,
    #     rightPadding=0 * cm,
    #     topPadding=1 * cm,
    #     showBoundary=0,
    #     id='text_frame')

    # # appending the report content to a list - pdf_content_list
    # report = Report.objects.get(id=pk)
    # report_content = ReportSection.objects.filter(report=report)

    # pdf_content_list = [Paragraph(report.name, stylesheet['Regular'])]
    # recorded_by = report.recorded_by.email
    # created_at = report.created_at
    # pdf_content_list.append(Paragraph(f"Recorded by: {recorded_by}", stylesheet['Regular']))
    # pdf_content_list.append(Paragraph(f"Date: {created_at}", stylesheet['Regular']))
    # if report.updated_at:
    #     updated_at = report.updated_at
    #     pdf_content_list.append(Paragraph(f"Updated at: {updated_at}", stylesheet['Regular']))
    # #decode to utf-8
    # for section in report_content.all():
    #     pdf_content_list.append(Paragraph(section.title, stylesheet['Bold']))
    #     pdf_content_list.append(Paragraph(section.section_content, stylesheet['Regular']))



    # # building the pdf_content
    # pdf_content_list.append(KeepTogether([]))
    
    # # temporary file location for the pdf
    # temporary_report_location = "temporary_report.pdf"
    
    # # establish a document
    # doc = BaseDocTemplate(temporary_report_location, pagesize=A4)

    # # creating a page template
    # frontpage = PageTemplate(id='FrontPage', frames=[text_frame])

    # # adding the pdf_content_list to the template and template to the document
    # doc.addPageTemplates(frontpage)

    # # building doc
    # doc.build(pdf_content_list)

    # """
    # Merging PDF files
    # """
    # current_report_document = report.report_document.path
    # merger = PdfMerger()
    
    # # append the files to the merger
    # merger.append(temporary_report_location)
    # merger.append(current_report_document)


    # # write out the merged PDF file on a temporary file location
    # temporary_merged_file_location = "temporary_merged_pages.pdf"
    # merger.write(temporary_merged_file_location)
    # merger.close()

    # # add watermark to merged PDF file
    # watermark_pdf =  './static/watermark.pdf'
    # watermarked_pdf_location = "watermarked_report.pdf"
    # watermark(content_pdf=temporary_merged_file_location, stamp_pdf=watermark_pdf, pdf_result=watermarked_pdf_location)

    # # convert to BytesIO object
    # with open(watermarked_pdf_location,'rb') as watermarked_pdf_file:
    #     watermarked_pdf_obj = io.BytesIO(watermarked_pdf_file.read())

    # # remove temporary files
    # os.remove(temporary_report_location)
    # os.remove(temporary_merged_file_location)
    # os.remove(watermarked_pdf_location)

    # return FileResponse(watermarked_pdf_obj, as_attachment=True, filename=f"{report.name}.pdf")



from django.shortcuts import render
from django.http import HttpResponse, HttpResponseRedirect
from django.urls import reverse
from docx import Document
from htmldocx import HtmlToDocx

from .models import Report, ReportSection
import pypandoc
from bs4 import BeautifulSoup
from docx import Document
from bleach.sanitizer import Cleaner
import logging
import html5lib, bs4
def validate_html_table(html_content):
    try:
        html5lib.HTMLParser().parse(html_content)
    except html5lib.html5parser.HTMLParseError as e:
        print("Invalid HTML structure:", e)
        print("Invalid HTML table structure in section content.")
        return False
    return True

def add_missing_tags(html_content):
    soup = bs4.BeautifulSoup(html_content, "html.parser")
    for table in soup.find_all("table"):
        for row in table.find_all("tr"):
            if len(row.find_all("td")) != len(row.find_all("th")):
                missing_cells = max(0, len(row.find_all("th")) - len(row.find_all("td")))
                for _ in range(missing_cells):
                    row.append(bs4.element.Tag("td"))
    return str(soup)

def check_nesting(html_content):
    soup = bs4.BeautifulSoup(html_content, "html.parser")
    for table in soup.find_all("table"):
        for row in table.find_all("tr"):
            for cell in row.find_all(["td", "th"]):
                if not cell.parent == row:
                    print("Incorrect nesting: Cell not within a row.")
                    
                    

ALLOWED_TAGS = [
    'p', 'br', 'a',
    'strong', 'em', 'i', 'b', 'u', 's',
    'blockquote', 'code',
    'table', 'thead', 'tbody', 'tr', 'td', 'ul', 'li',
    'p/',  # include "p" with closing slash
]

ALLOWED_ATTRS = {
    '*': ['class', 'style'],  # Allow basic styling
    'a': ['href', 'title'],
}

# Create cleaner with extended tags
cleaner = Cleaner(
    tags=ALLOWED_TAGS,
    attributes=ALLOWED_ATTRS,
    strip=True,  # Remove disallowed tags and attributes
)


# def pydoc(html):
#     try:
#         temporary_report_location = "temporary_report.pdf"
#         output = pypandoc.convert_file(html, 'pdf', outputfile=temporary_report_location,
#         extra_args=['-V', 'geometry:margin=1.5cm'])
#         response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#         response['Content-Disposition'] = 'attachment; filename="exported_content.docx"'
#         return response
#     except Exception as e:
#         # Handle conversion errors with PyPandoc
#         print(f"PyPandoc conversion error: {e}")
#         return 0


from docx.shared import Inches  

def handle_table(table_element, document):
    table_rows = []
    for row in table_element.find_all('tr'):
        table_cells = []
        for cell in row.find_all(['th', 'td']):
            cell_text = cell.text.strip()  # Extract text content
            cell_span = int(cell.get('colspan', 1))  # Handle column spans
            cell_rows = int(cell.get('rowspan', 1))  # Handle row spans

            # Logic for handling nested headers and column spans (implementation details will depend on your table structure)

            table_cells.extend([cell_text] * cell_span)  # Account for column spans

        table_rows.append(table_cells)

    # Adjust table structure based on nested headers and column spans (implementation details will depend on your requirements)

    # Create a Word table with appropriate formatting
    table = document.add_table(rows=len(table_rows), cols=max(len(row) for row in table_rows))
    for i, row in enumerate(table_rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text

    return table

                            
                                   
def generate_report_word_document(request, report_id):
    try:
        report = Report.objects.get(id=report_id)
        sections = ReportSection.objects.filter(report=report)
    except Report.DoesNotExist:  # Use specific model exception for clarity
        return HttpResponse("Report not found.", status=404)

    document = Document()
    new_parser = HtmlToDocx()
    new_parser.table_style = 'Light Shading Accent 4'
    new_parser.table_autoformat=True
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "Master Report"
    
    

    document.add_heading(report.name, 0)
    # Create content directly within the document (more efficient than HTML conversion)
    for section in sections:
        document.add_heading(section.title, level=1)
        if section.section_content:


            # Create dictionaries to store list and image data (if applicable)
            lists = {}  # Track list IDs and items
            images = []  # Store image data for later insertion

            # Parse the HTML content using BeautifulSoup
            soup = BeautifulSoup(section.section_content, 'html.parser')

            # Extract text, headings, paragraphs, lists, images, and tables
            for element in soup.find_all(True, recursive=False):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # Add headings with appropriate styles
                    paragraph = document.add_paragraph(element.text, style='Heading ' + element.name[1])
                elif element.name == 'p':
                    paragraph = document.add_paragraph(element.text)
                elif element.name == 'ul' or element.name == 'ol':
                    list_id = element.get('id', None)  # Assign unique ID if not present
                    if list_id:
                        lists[list_id] = []
                    else:
                        print("Warning: List without ID encountered. Using sequential numbering.")
                        lists[len(lists)] = []
                    for item in element.find_all('li'):
                        lists[list_id].append(item.text)
                elif element.name == 'img':
                    image_url = element.get('src')
                    images.append((image_url, element.attrs.get('style', None)))  # Store URL and style
                elif element.name == 'table':
                    # Table handling logic (see below for details)
                    table = handle_table(element, document)
                    if table:
                        pass

            # Create lists in the Word document based on extracted data
            for list_id, items in lists.items():
                paragraph = document.add_paragraph()
                if list_id.isdigit():
                    list_type = 'NUMBER'
                else:
                    list_type = 'BULLET' if element.name == 'ul' else 'ARABIC'
                list_format = document.lists[0].bullet_format  # Create a new list format
                for level in [0, 1, 2, 3, 4]:  # Create levels up to 4 (adjust as needed)
                    list_format.create_level(level=level, indent=Inches(0.25 * level), text="{}. ".format(str(level + 1)),
                                            font=list_format.font)
                paragraph.add_run("** {} **".format(list_id)).bold = True
                for item in items:
                    paragraph.add_run("   ").bullet_format = list_format

            # Insert images at their original positions (adapt based on your HTML structure)
            for image_url, image_style in images:
                paragraph = document.paragraphs[-1]  # Insert after the last processed element
                if image_url:
                    try:
                        paragraph.add_picture(image_url, width=Inches(4))  # Adjust width as needed
                    except Exception as e:
                        print("Error inserting image:", e)
                    if image_style:
                        # Apply image style attributes (adapt based on desired properties)
                        run = paragraph.runs[-1]
                        for attr, value in image_style.items():
                            setattr(run, attr, value)

   

    # Generate unique filename with timestamp for clarity
    filename = f"{report.name}-{report_id}-{datetime.date.today().strftime('%Y%m%d_%H%M%S')}.docx"

    # Save directly to a temporary file using a context manager
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        output_path = temp_file.name
        document.save(output_path)

    # Read the temporary file and return as download response
    with open(output_path, 'rb') as file:
        file_content = file.read()
    response = HttpResponse(file_content, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f"attachment; filename={filename}"
    return response





def export_report_to_word(request, report_id):
    document = export_to_word(request, report_id)
    if document.status_code == 200:
        return document
    else:
        return HttpResponse("An error occurred during report generation.", status=500)

from docx.shared import Inches
from docx import Document
from bs4 import BeautifulSoup
# def export_to_word(request, content_id):
#     """Exports specified content as a Word document."""
#     report = Report.objects.get(id=content_id)
#     content = ReportSection.objects.filter(report=report).first()

#     html_content = content.section_content

#     doc = Document()  # Create a new Word document
#     tables = {}  # Store table data

#     soup = BeautifulSoup(html_content, 'html.parser')

#     def handle_element(element):
#         """Handles different HTML elements and adds them to the Word document."""

#         if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
#             doc.add_paragraph(element.text, style='Heading ' + element.name[1])
#         elif element.name == 'p':
#             doc.add_paragraph(element.text)
#         elif element.name == 'table':
#             handle_table(element)
#         # elif element.name in ['ul', 'ol']:
#         #     handle_list(element)
#         elif element.name in ['div', 'blockquote']:
#             handle_block_element(element)
#         elif element.name == 'img':
#             handle_image(element)

#     def handle_table(table_element):
#         """Extracts table data and adds it to the Word document."""

#         table_id = table_element.get('id', None)
#         if not table_id:
#             warnings.warn("Table without ID encountered. Using sequential numbering.")
#             table_id = len(tables)

#         tables[table_id] = []
#         for row in table_element.find_all('tr'):
#             table_row = []
#             for cell in row.find_all('td', 'th'):
#                 table_row.append(cell.text)
#             tables[table_id].append(table_row)

#     def handle_list(list_element):
#         """Creates a list in the Word document."""

#         list_obj = doc.add_paragraph()
#         list_format = ListFormat.BULLET if list_element.name == 'ul' else ListFormat.NUMBERED

#         for item in list_element.find_all(['li']):
#             list_obj.add_run(item.text).font.name = 'Symbol'  # Use Symbol font for bullets
#             list_obj.paragraph_format.bullet.list_format = list_format

#     def handle_block_element(element):
#         """Adds styles or handles block elements specifically."""

#         paragraph = doc.add_paragraph(element.text)
#         if element.name == 'blockquote':
#             paragraph.paragraph_format.left_indent = Inches(0.5)

#     def handle_image(image_element):
#         """Adds images to the Word document."""

#         image_path = image_element.get('src')
#         try:
#             doc.add_picture(image_path, width=Inches(3.0))  # Adjust width as needed
#         except FileNotFoundError:
#             warnings.warn(f"Warning: Image '{image_path}' not found.")

#     # Process all top-level HTML elements
#     for element in soup.find_all(True, recursive=False):
#         handle_element(element)

#     # Create tables in the Word document
#     for table_id, table_data in tables.items():
#         table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
#         for i, row in enumerate(table_data):
#             for j, cell_text in enumerate(row):
#                 cell = table.cell(i, j)
#                 cell.text = cell_text

           
#     doc.save('exported_content.docx')
#     with open('exported_content.docx', 'rb') as f:
#         response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#         response['Content-Disposition'] = 'attachment; filename="exported_content.docx"'
#         return response       


def pdf_report(request, pk):
    """
    Generating the PDF file for the report content
    """
    #Stylesheet configs
    font_family = 'Amharic'
    font_size = 10
    document = generate_report_word_document(request, pk)
    return document    
    
def viewMasterReport(request):
    user_sector = UserSector.objects.get(user=request.user.id).user_sector

    # Get the report type with a deadline greater than today
    report_type = ReportType.objects.filter(deadline__gt=date.today()).first()

    # Get the master reports for the user's ministry
    master_reports = MasterReport.objects.filter(responsible_ministry=user_sector)

    # Get the current year
    current_year = date.today().year +2 

    # Get the years that are less than or equal to the current year and are visible
    years = Year.objects.filter(year_eng__lte=current_year, visible=True)

    # Get all quarters
    quarters = Quarter.objects.all().order_by('id')

    # Get the current master report for the user's ministry and the selected report type
    current = master_reports.filter(report_type=report_type).first()
    

    if request.method == 'POST':
        file = request.FILES['file']
        if current is not None and report_type is not None:
            current.report_doc = file
            current.save()
        elif report_type is not None and current is None:
            model_instance = MasterReport()  
            model_instance.report_doc = file
            model_instance.year = Year.objects.get(id = report_type.year.id)
            model_instance.quarter = Quarter.objects.get(id = report_type.quarter.id)
            model_instance.report_type = report_type
            model_instance.responsible_ministry = UserSector.objects.get(user=request.user.id).user_sector
            model_instance.save()

        # Redirect to a success page or render a template
        return redirect('viewMasterReport') 

    context = {
        'report_type': report_type,
        'years': years,
        'quarters': quarters,
        'master_reports': master_reports,
        'current': current,
    }

    return render(request, 'view_master_report.html', context)




from datetime import datetime

def view_mopd_report(request):
    # Get the selected year from the request
    selected_year = request.GET.getlist('selected_year[]')
    
    # Fetch the latest report type object
    obj = ReportType.objects.order_by('-created_at').first()
    ministry = ResponsibleMinistry.objects.filter(ministry_is_visable =True)
    
    if not selected_year:
        # If no year is selected, filter by the year marked as the current year
        current_year_obj = Year.objects.filter(is_current_year=True).first()
        if current_year_obj:
            master_reports = MasterReport.objects.filter(year=current_year_obj.year_eng)
        else:
            # Fallback in case no year is marked as the current year
            current_year = datetime.now().year
            master_reports = MasterReport.objects.filter(year=current_year)
    else:
        # If a year is selected, filter by the selected year
        master_reports = MasterReport.objects.filter(year__id=int(selected_year[0]))

    quarters = Quarter.objects.all().order_by('id')
    years = Year.objects.all()

    return render(request, 'view_mopd_master_report.html', {
        'master_reports': master_reports,
        'ministry': ministry,
        'obj': obj,
        'quarters': quarters,
        'years': years
    })





def view_mopd_report_active(request):
    obj=  ReportType.objects.filter(deadline__gte= date.today()).first()
    ministry = ResponsibleMinistry.objects.all()
    master_reports = MasterReport.objects.filter(report_type__id = obj.id)
    
    
    return render(request, 'view_mopd_report_active.html', {'master_reports': master_reports,'ministry':ministry,'obj':obj})